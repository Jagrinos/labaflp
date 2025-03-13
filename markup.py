import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Mm, Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def setupRun(r_fmt):
    r_fmt.name = 'Times new roman'
    r_fmt.size = Pt(14)

def markupdocx(recipient, title, text, sender_profession, sender):
    doc = Document()

    #picture
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture('test.jpg', width=Mm(25))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #company
    p = doc.add_paragraph()
    run = p.add_run("ООО «РУССКИЙ СТИЛЬ – 97»")
    setupRun(run.font)
    run.font.bold = True

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)

    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc.get_or_add_tcPr().append(docx.oxml.OxmlElement('w:tblBorders'))

    cell = table.cell(0, 0)
    cell._element.clear()
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = Pt(15)



    run = p.add_run( "ООО \"Русский Стиль - 97\"\n"
                    "Место нахождения: Симферопольская ул., д. 55 Краснодар, 350080\n"
                    "Почтовый адрес: Симферопольская ул., д. 55\n"
                    "Краснодар, 350080\n"
                    "Тел. /факс: (861) 260-09-22 , 260-09-23\n"
                    "ОГРН 1022301977400\n"
                    "ИНН/КПП 2312065374/231201001\n"
                     "\n"
                     "ИСХ № ________от _________    ")
    setupRun(run.font)

    cell = table.cell(0, 1)
    cell._element.clear()
    p = cell.add_paragraph(recipient)
    p.paragraph_format.line_spacing = Pt(15)
    run = p.add_run()
    setupRun(run.font)

    #start text
    p = doc.add_paragraph()
    run = p.add_run("\n"+title)
    setupRun(run.font)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraphs = text.split("\n")
    for ptx in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(1.0)
        p.paragraph_format.line_spacing = Pt(15)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(ptx)
        setupRun(run.font)


    #table last
    p = doc.add_paragraph()
    p = doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)

    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc.get_or_add_tcPr().append(docx.oxml.OxmlElement('w:tblBorders'))

    cell = table.cell(0, 0)
    cell._element.clear()
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = Pt(15)
    run = p.add_run(sender_profession)
    setupRun(run.font)

    cell = table.cell(0, 1)
    cell._element.clear()
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = Pt(15)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run(sender)
    setupRun(run.font)

    doc.save("output.docx")

